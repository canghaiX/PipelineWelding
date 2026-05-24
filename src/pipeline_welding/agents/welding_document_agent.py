from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class WeldingDocumentAgentConfig:
    template_docx_path: Path = Path("data/MHPWPS-062.docx")
    output_dir: Path = Path("result")
    output_prefix: str = "welding_standard_filled"


class WeldingDocumentAgent:
    """Fills a WPS-style DOCX from the standard-agent JSON result."""

    def __init__(self, config: WeldingDocumentAgentConfig | None = None) -> None:
        self.config = config or WeldingDocumentAgentConfig()

    def build_document(self, standard_result: dict[str, Any]) -> Path:
        template_path = self.config.template_docx_path
        if not template_path.exists():
            raise FileNotFoundError(f"Template DOCX not found: {template_path}")

        self.config.output_dir.mkdir(parents=True, exist_ok=True)
        document = Document(str(template_path))

        document_fields = self._extract_document_fields(standard_result)
        self._fill_known_paragraphs(document, document_fields)
        self._fill_known_tables(document, document_fields)

        output_path = self._build_output_path()
        document.save(str(output_path))
        print(f"已生成焊接标准文档：{output_path}")
        return output_path

    def _extract_document_fields(self, standard_result: dict[str, Any]) -> dict[str, str]:
        document_fields = standard_result.get("document_fields", {})
        if isinstance(document_fields, dict) and document_fields:
            return {key: self._clean_value(value) for key, value in document_fields.items()}

        fields = standard_result.get("input", {})
        if not isinstance(fields, dict):
            return {}
        fallback = {key: self._clean_value(value) for key, value in fields.items()}
        fallback.update(self._extract_reference_values(standard_result))
        return fallback

    def _fill_known_paragraphs(
        self,
        document: Document,
        document_fields: dict[str, str],
    ) -> None:
        replacements: dict[str, str] = self._paragraph_replacements(document_fields)

        for paragraph in document.paragraphs:
            self._replace_paragraph_by_keyword(paragraph, replacements)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_paragraph_by_keyword(paragraph, replacements)

    def _paragraph_replacements(self, document_fields: dict[str, str]) -> dict[str, str]:
        material = self._value(document_fields, "base_material")
        size = self._value(document_fields, "base_thickness_or_diameter")
        return {
            "焊接方法": self._line_value("焊接方法", self._value(document_fields, "welding_process")),
            "坡口形式": self._line_value("坡口形式:", self._value(document_fields, "joint_type")),
            "类别号": f"母材：标准号/材料代号 {material}",
            "对接焊缝焊件母材厚度范围": self._line_value("对接焊缝焊件母材厚度范围", size),
            "管子直径、壁厚范围": self._line_value("管子直径、壁厚范围：对接焊缝", size),
            "最小预热温度": self._line_value("最小预热温度（℃）", self._value(document_fields, "preheat_temperature")),
            "最大道间温度": self._line_value("最大道间温度（℃）", self._value(document_fields, "interpass_temperature")),
        }

    def _fill_known_tables(
        self,
        document: Document,
        document_fields: dict[str, str],
    ) -> None:
        for index, table in enumerate(document.tables):
            if index == 0:
                self._fill_joint_table(table, document_fields)
            if index == 1:
                self._fill_filler_metal_table(table, document_fields)
            if index == 3:
                self._fill_process_condition_table(table, document_fields)
            self._fill_labeled_table(table, document_fields)
            if self._is_welding_parameter_table(table):
                self._fill_welding_parameter_table(table, document_fields)

    def _fill_joint_table(self, table: Any, document_fields: dict[str, str]) -> None:
        if not table.rows or not table.rows[0].cells:
            return
        table.rows[0].cells[0].text = (
            "焊接接头："
            f"坡口形式:{self._value(document_fields, 'joint_type')}"
            "衬垫(材料及规格)/"
            "其他/"
        )

    def _fill_filler_metal_table(self, table: Any, document_fields: dict[str, str]) -> None:
        row_values = {
            "焊材标准": self._value(document_fields, "filler_standard"),
            "填充金属尺寸": self._value(document_fields, "filler_diameter"),
            "焊材型号": self._value(document_fields, "filler_metal"),
            "焊材牌号": self._value(document_fields, "filler_metal"),
            "填充金属类别": self._value(document_fields, "filler_category"),
        }
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip().replace("：", "")
            for key, value in row_values.items():
                if key in label:
                    row.cells[1].text = value
                    if len(row.cells) > 2:
                        row.cells[2].text = "/"
                    break

    def _fill_process_condition_table(self, table: Any, document_fields: dict[str, str]) -> None:
        if len(table.rows) < 3:
            return
        table.rows[1].cells[0].text = (
            "预热："
            f"最小预热温度（℃）{self._value(document_fields, 'preheat_temperature')}"
            f"最大道间温度（℃）{self._value(document_fields, 'interpass_temperature')}"
            "保持预热时间/"
            "加热方式/"
        )
        table.rows[1].cells[1].text = (
            "气体："
            "气体种类 混合比 流量(L/min)"
            f"保护气 {self._value(document_fields, 'shielding_gas')} / {self._value(document_fields, 'gas_flow')}"
            "尾部保护气 / / /"
            "背面保护气 / / /"
        )

    def _fill_labeled_table(self, table: Any, document_fields: dict[str, str]) -> None:
        label_to_value = {
            "填充金属尺寸": document_fields.get("filler_diameter"),
            "焊材型号": document_fields.get("filler_metal"),
            "焊材牌号": document_fields.get("filler_metal"),
        }
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip().replace("：", "")
            for expected_label, value in label_to_value.items():
                if value and expected_label in label:
                    row.cells[1].text = value
                    break

    @staticmethod
    def _is_welding_parameter_table(table: Any) -> bool:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        return all(keyword in text for keyword in ("焊道/焊层", "焊接方法", "焊接电流"))

    def _fill_welding_parameter_table(
        self,
        table: Any,
        document_fields: dict[str, str],
    ) -> None:
        welding_process = self._value(document_fields, "welding_process")
        filler_metal = self._value(document_fields, "filler_metal")
        filler_diameter = self._value(document_fields, "filler_diameter")
        polarity = self._value(document_fields, "polarity")
        current = self._value(document_fields, "current")
        voltage = self._value(document_fields, "voltage")
        speed = self._value(document_fields, "welding_speed")
        heat_input = self._value(document_fields, "heat_input")

        for row in table.rows:
            cells = row.cells
            if not cells or not cells[0].text.strip().isdigit():
                continue
            self._set_cell_text(cells, 1, welding_process)
            self._set_cell_text(cells, 2, filler_metal)
            self._set_cell_text(cells, 3, filler_diameter)
            self._set_cell_text(cells, 4, polarity)
            self._set_cell_text(cells, 5, current)
            self._set_cell_text(cells, 6, voltage)
            self._set_cell_text(cells, 7, speed)
            self._set_cell_text(cells, 8, heat_input)

    @staticmethod
    def _line_value(label: str, value: str | None) -> str:
        return f"{label}{WeldingDocumentAgent._clean_value(value)}"

    @staticmethod
    def _base_material_line(value: str | None) -> str:
        return f"母材：标准号/材料代号 {WeldingDocumentAgent._clean_value(value)}"

    def _extract_reference_values(self, standard_result: dict[str, Any]) -> dict[str, str]:
        text = self._collect_clean_reference_text(standard_result)
        return {
            "preheat_temperature": self._first_match(text, (r"预热温度[^\d≥≤]*([≥≤]?\s*\d+\s*℃?)",)),
            "interpass_temperature": self._first_match(text, (r"层间温度[^\d≥≤]*([≥≤]?\s*\d+\s*℃?)", r"道间温度[^\d≥≤]*([≥≤]?\s*\d+\s*℃?)")),
            "current": self._first_match(text, (r"电流[^\d]*(\d+\s*[-~～]\s*\d+\s*A?)",)),
            "voltage": self._first_match(text, (r"电压[^\d]*(\d+\s*[-~～]\s*\d+\s*V?)",)),
            "welding_speed": self._first_match(text, (r"焊接速度[^\d]*(\d+\s*[-~～]\s*\d+\s*(?:mm/min)?)",)),
            "heat_input": self._first_match(text, (r"线能量[^\d]*(\d+(?:\.\d+)?\s*(?:kJ/cm)?)",)),
            "filler_metal": self._first_match(text, (r"(E\d{3,4}[A-Z0-9-]*)",)),
            "filler_diameter": self._first_match(text, (r"(?:φ|Φ)\s*(\d+(?:\.\d+)?\s*mm)",)),
            "polarity": self._first_match(text, (r"(反接|正接|DCEN|DCEP|EP|EN)",)),
        }

    def _collect_clean_reference_text(self, standard_result: dict[str, Any]) -> str:
        parts: list[str] = []
        standard = standard_result.get("pipeline_welding_standard", {})
        if isinstance(standard, dict):
            parts.extend(str(item) for item in standard.get("required_controls", []))

        results = standard_result.get("mcp_search", {}).get("results", {})
        if isinstance(results, dict):
            for items in results.values():
                if isinstance(items, list):
                    for item in items:
                        if isinstance(item, dict):
                            parts.extend(
                                str(item.get(key, "")) for key in ("title", "snippet") if item.get(key)
                            )

        return "\n".join(part for part in parts if self._is_clean_text(part))

    @staticmethod
    def _is_clean_text(text: str) -> bool:
        lowered = text.lower()
        blocked = ("not found", "unknown tool", "unknown too", "error", "missing")
        return bool(text.strip()) and not any(item in lowered for item in blocked)

    @staticmethod
    def _first_match(text: str, patterns: tuple[str, ...]) -> str:
        import re

        for pattern in patterns:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                return match.group(1).replace(" ", "")
        return ""

    def _build_output_path(self) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return self.config.output_dir / f"{self.config.output_prefix}_{timestamp}.docx"

    @staticmethod
    def _replace_paragraph_by_keyword(paragraph: Paragraph, replacements: dict[str, str]) -> None:
        text = paragraph.text.strip()
        for keyword, value in replacements.items():
            if keyword in text:
                paragraph.text = value
                return

    @staticmethod
    def _set_cell_text(cells: Any, index: int, value: str | None) -> None:
        if index < len(cells):
            cells[index].text = WeldingDocumentAgent._clean_value(value)

    @staticmethod
    def _clean_value(value: Any) -> str:
        text = str(value).strip() if value is not None else ""
        lowered = text.lower()
        blocked = ("not found", "unknown tool", "unknown too", "error", "missing")
        if not text or any(item in lowered for item in blocked):
            return "/"
        return text

    @staticmethod
    def _value(document_fields: dict[str, str], key: str) -> str:
        return WeldingDocumentAgent._clean_value(document_fields.get(key))


def build_welding_document_agent(
    config: WeldingDocumentAgentConfig | None = None,
) -> WeldingDocumentAgent:
    return WeldingDocumentAgent(config=config)


def build_welding_document_agent_from_config(config: dict[str, Any]) -> WeldingDocumentAgent:
    template_config = config.get("template", {})
    output_config = config.get("output", {})
    return WeldingDocumentAgent(
        WeldingDocumentAgentConfig(
            template_docx_path=Path(template_config.get("docx_path", "data/MHPWPS-062.docx")),
            output_dir=Path(output_config.get("dir", "result")),
            output_prefix=output_config.get("prefix", "welding_standard_filled"),
        )
    )
